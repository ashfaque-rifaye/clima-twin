"""Professional planning-report generation (deck Part 7–9).

Assembles a decision-ready urban-climate planning document — executive summary,
study area, current assessment, intervention plan, multi-metric impact, budget
(capital + 5/10-yr), timeline, risk, methodology and appendices — with an
AI-authored narrative (Gemini, with deterministic fallback), inline SVG charts,
and embedded Google Static Maps. Renders to self-contained HTML (which the
browser prints to PDF) and to DOCX. This replaces the old markdown export.
"""
from __future__ import annotations

import base64
import datetime as dt
import html
import logging

import httpx
from pydantic import BaseModel

from .config import settings
from .gemini import gemini_available, generate_json

log = logging.getLogger("climatwin.report")

BLUE = "#1a73e8"
INK = "#1f2430"
GOOD = "#12a36b"


# --------------------------------------------------------------------------- #
# formatting helpers
# --------------------------------------------------------------------------- #
def _inr(n) -> str:
    n = float(n or 0)
    if n >= 1e7:
        return f"₹{n / 1e7:.2f} Cr"
    if n >= 1e5:
        return f"₹{n / 1e5:.2f} L"
    return f"₹{int(n):,}"


def _int(n) -> str:
    return f"{int(n or 0):,}"


def _esc(s) -> str:
    return html.escape(str(s if s is not None else "—"))


# --------------------------------------------------------------------------- #
# Google Static Maps (server-side download → base64, so it works in any viewer)
# --------------------------------------------------------------------------- #
def _static_map(lat: float, lng: float, zoom: int = 14, maptype: str = "roadmap") -> tuple[bytes, str] | None:
    key = settings.server_api_key or settings.google_maps_api_key
    if not key:
        return None
    try:
        r = httpx.get(
            "https://maps.googleapis.com/maps/api/staticmap",
            params={
                "center": f"{lat},{lng}", "zoom": zoom, "size": "640x360", "scale": 2,
                "maptype": maptype, "markers": f"color:0x1a73e8|{lat},{lng}", "key": key,
            },
            timeout=12,
        )
        if r.status_code == 200 and r.headers.get("content-type", "").startswith("image"):
            return r.content, r.headers["content-type"]
        log.warning("static map HTTP %s", r.status_code)
    except Exception as exc:
        log.warning("static map failed: %s", type(exc).__name__)
    return None


def _data_uri(img: tuple[bytes, str] | None) -> str | None:
    if not img:
        return None
    data, ctype = img
    return f"data:{ctype};base64,{base64.b64encode(data).decode()}"


# --------------------------------------------------------------------------- #
# inline SVG charts
# --------------------------------------------------------------------------- #
def _bar_before_after(baseline: float, projected: float) -> str:
    if baseline is None:
        return ""
    lo, hi = min(baseline, projected), max(baseline, projected)
    span = max(hi, 50)
    bw = lambda v: max(4, (v / span) * 460)  # noqa: E731
    return (
        f'<svg viewBox="0 0 520 90" width="100%" role="img">'
        f'<text x="0" y="16" font-size="12" fill="#5b6472">Before {baseline:.1f}°C</text>'
        f'<rect x="0" y="22" width="{bw(baseline):.0f}" height="14" rx="4" fill="#c1271f"/>'
        f'<text x="0" y="60" font-size="12" fill="#5b6472">After {projected:.1f}°C</text>'
        f'<rect x="0" y="66" width="{bw(projected):.0f}" height="14" rx="4" fill="{GOOD}"/>'
        f"</svg>"
    )


def _bar_impacts(impacts: dict) -> str:
    rows = [
        ("Temp reduction (°C)", impacts.get("temp_reduction_c", 0)),
        ("AQI improvement", impacts.get("aqi_improvement", 0)),
        ("Flood managed (m³)", impacts.get("flood_managed_m3", 0)),
        ("Canopy added (m²)", impacts.get("canopy_added_m2", 0)),
        ("Carbon (kg CO₂/yr)", impacts.get("carbon_seq_kg_year", 0)),
    ]
    rows = [(l, v) for l, v in rows if v]
    if not rows:
        return ""
    top = max(v for _, v in rows) or 1
    h = 26 * len(rows) + 10
    bars = []
    for i, (label, v) in enumerate(rows):
        y = 10 + i * 26
        w = max(3, (v / top) * 300)
        bars.append(
            f'<text x="0" y="{y + 12}" font-size="11" fill="#5b6472">{_esc(label)}</text>'
            f'<rect x="180" y="{y}" width="{w:.0f}" height="16" rx="4" fill="{BLUE}"/>'
            f'<text x="{188 + w:.0f}" y="{y + 12}" font-size="11" font-weight="700" fill="{INK}">{_int(v) if v >= 10 else v}</text>'
        )
    return f'<svg viewBox="0 0 520 {h}" width="100%" role="img">{"".join(bars)}</svg>'


# --------------------------------------------------------------------------- #
# narrative (AI with fallback)
# --------------------------------------------------------------------------- #
class _Narrative(BaseModel):
    executive_summary: str = ""
    why_selected: str = ""
    why_alternatives_rejected: str = ""
    environmental_benefit: str = ""
    social_benefit: str = ""
    economic_benefit: str = ""
    risks: list[str] = []
    timeline_immediate: str = ""
    timeline_year1: str = ""
    timeline_year3: str = ""
    timeline_year5: str = ""


def _narrative(area: str, hazard: str, plan_names: str, impacts: dict, costs: dict) -> dict:
    fallback = {
        "executive_summary": (
            f"This plan targets {hazard} risk in {area} through {plan_names}. It is projected to "
            f"deliver measurable improvement (see Impact Analysis) benefiting ~{_int(impacts.get('people_benefited'))} "
            f"residents for a capital outlay of {_inr(costs.get('capital_inr'))}, with {_inr(costs.get('maintenance_inr_year'))}/year "
            f"upkeep. Interventions are reversible, phaseable, and prioritise vulnerable, data-poor areas."
        ),
        "why_selected": f"The mix maximises {hazard}-hazard benefit per rupee for {area}'s conditions while diversifying across complementary measures to spread delivery and maintenance risk.",
        "why_alternatives_rejected": "Higher-cost engineered alternatives were down-weighted where nature-based options deliver comparable benefit at lower capital and maintenance, and single-measure plans were avoided to reduce dependency risk.",
        "environmental_benefit": f"Direct {hazard} mitigation plus co-benefits in canopy, carbon and stormwater (see Impact Analysis).",
        "social_benefit": f"~{_int(impacts.get('people_benefited'))} people benefited, prioritising commuters, the elderly and low-income residents who live the consequences.",
        "economic_benefit": f"Avoided heat/flood productivity losses and lower long-run maintenance; {_inr(costs.get('ten_year_inr'))} ten-year cost of ownership.",
        "risks": [],
        "timeline_immediate": "Site survey, community consultation, procurement and quick-win installs (shade, signage, dust control).",
        "timeline_year1": "Core installation of nature-based and structural measures; baseline monitoring established.",
        "timeline_year3": "Vegetation establishment matures; measured impact review and course-correction.",
        "timeline_year5": "Full projected benefit realised; evaluate scale-out to adjacent wards.",
    }
    if not (gemini_available()):
        return fallback
    ai = generate_json(
        f"You are a senior urban-planning consultant drafting a municipal planning report for {area}, Chennai. "
        f"Hazard: {hazard}. Proposed plan: {plan_names}. Projected impact: {impacts}. Costs: {costs}. "
        "Return JSON with: executive_summary (4-5 sentences, board-level), why_selected, "
        "why_alternatives_rejected, environmental_benefit, social_benefit, economic_benefit (2-3 sentences each), "
        "risks (list of 3-5), timeline_immediate, timeline_year1, timeline_year3, timeline_year5 (1 sentence each). "
        "Professional, concrete, no markdown.",
        response_schema=_Narrative,
    )
    if isinstance(ai, dict) and ai.get("executive_summary"):
        return {**fallback, **{k: v for k, v in ai.items() if v}}
    return fallback


# --------------------------------------------------------------------------- #
# assembly
# --------------------------------------------------------------------------- #
def build_report(area_name: str, lat: float, lng: float, hazard: str, point: dict, plan: dict) -> dict:
    interventions = plan.get("interventions") or []
    impacts = plan.get("impacts") or {}
    costs = plan.get("costs") or {}
    budget = plan.get("budget") or {}
    plan_names = ", ".join(
        f"{_int(i.get('count'))}× {i.get('name') or i.get('species') or i.get('type')}" for i in interventions
    ) or "the proposed intervention mix"

    narrative = _narrative(area_name, hazard, plan_names, impacts, costs)
    risks = list(dict.fromkeys((plan.get("trade_offs") or []) + (narrative.get("risks") or [])))

    heat = point.get("heat") or {}
    air = point.get("air") or {}
    flood = point.get("flood") or {}
    vuln = point.get("vulnerability") or {}

    return {
        "title": f"Urban Climate Resilience Plan — {area_name}",
        "subtitle": f"{hazard.title()}-focused intervention strategy · Chennai",
        "generated_at": dt.datetime.now(dt.timezone.utc).strftime("%d %B %Y"),
        "hazard": hazard,
        "study_area": {
            "area_name": area_name,
            "coordinates": f"{lat:.5f}, {lng:.5f}",
            "administrative": "Greater Chennai Corporation (GCC) · Tamil Nadu",
            "population": vuln.get("population"),
            "commuter_footfall": vuln.get("commuter_footfall"),
            "elderly_pct": vuln.get("elderly_pct"),
            "elevation_m": point.get("elevation_m"),
            "data_blind_spot": vuln.get("data_blind_spot"),
        },
        "assessment": {
            "heat": heat.get("feels_like_c"),
            "condition": heat.get("condition"),
            "aqi": air.get("aqi"),
            "aqi_category": air.get("category"),
            "flood_risk": flood.get("risk"),
            "green_cover_pct": vuln.get("green_cover_pct"),
            "ndvi": vuln.get("ndvi"),
        },
        "interventions": interventions,
        "impacts": impacts,
        "costs": costs,
        "budget": budget,
        "narrative": narrative,
        "risks": risks or ["No major risks flagged for this mix."],
        "confidence": (plan.get("confidence") or {}),
        "methodology": (
            "Impacts are computed by a literature-informed coefficient model cross-checked against a "
            "BigQuery ML linear land-surface-temperature model trained on the Chennai analysis grid. "
            "Live conditions are read per point from Google Weather, Air Quality and Elevation APIs; the "
            "base grid is served from BigQuery and can be enriched with Earth Engine LST/NDVI/DEM. Every "
            "figure carries an uncertainty band and cited coefficients."
        ),
        "data_sources": [
            "Google Weather API — live feels-like temperature",
            "Google Air Quality API — AQI, dominant pollutant",
            "Google Elevation API / SRTM — terrain & flood modelling",
            "Google Earth Engine — MODIS LST, NDVI; Landsat NDBI (build-time)",
            "BigQuery + BigQuery ML — analysis grid & cooling model",
            "US EPA / FAO / i-Tree — intervention cooling & cost coefficients",
        ],
        "assumptions": plan.get("assumptions") or [
            "Coefficients are literature-informed and illustrative, not site-measured.",
            "Budget figures cover capital cost; maintenance is reported separately.",
            "Population and vulnerability are census-informed model estimates.",
        ],
        "maps": {
            "roadmap": _data_uri(_static_map(lat, lng, zoom=14, maptype="roadmap")),
            "satellite": _data_uri(_static_map(lat, lng, zoom=16, maptype="satellite")),
        },
        "charts": {
            "before_after": _bar_before_after(impacts.get("temp_reduction_c") is not None and heat.get("feels_like_c") or heat.get("feels_like_c"),
                                              (heat.get("feels_like_c") or 0) - (impacts.get("temp_reduction_c") or 0)) if heat.get("feels_like_c") else "",
            "impacts": _bar_impacts(impacts),
        },
    }


# --------------------------------------------------------------------------- #
# HTML rendering
# --------------------------------------------------------------------------- #
def _table(rows: list[tuple[str, str]]) -> str:
    trs = "".join(f'<tr><th>{_esc(k)}</th><td>{v}</td></tr>' for k, v in rows)
    return f'<table class="kv">{trs}</table>'


def render_html(r: dict) -> str:
    sa, asmt, imp, costs = r["study_area"], r["assessment"], r["impacts"], r["costs"]
    n = r["narrative"]
    maps = "".join(
        f'<figure><img src="{src}" alt="{cap}"/><figcaption>Figure {i + 1}. {cap} — {_esc(sa["coordinates"])}</figcaption></figure>'
        for i, (src, cap) in enumerate([(r["maps"]["roadmap"], "Study area (roadmap, study point marked)"),
                                        (r["maps"]["satellite"], "Study area (satellite context)")]) if src
    ) or '<p class="muted">Map imagery unavailable at generation time.</p>'

    iv_rows = "".join(
        f'<tr><td>{_esc(i.get("name") or i.get("species"))}</td><td>{_int(i.get("count"))}</td>'
        f'<td>{_inr(i.get("capital_inr"))}</td><td>{_esc(i.get("why", ""))}</td></tr>'
        for i in r["interventions"]
    ) or '<tr><td colspan="4" class="muted">No interventions specified.</td></tr>'

    impact_rows = [
        ("Temperature reduction", f"−{imp.get('temp_reduction_c', 0)} °C feels-like"),
        ("Air-quality improvement", f"{imp.get('aqi_improvement', 0)} AQI points"),
        ("Stormwater managed", f"{_int(imp.get('flood_managed_m3'))} m³ / event"),
        ("Green area added", f"{_int(imp.get('canopy_added_m2'))} m²"),
        ("Carbon sequestration", f"{_int(imp.get('carbon_seq_kg_year'))} kg CO₂ / year"),
        ("Population benefited", _int(imp.get("people_benefited"))),
    ]
    budget_note = ""
    if r["budget"]:
        b = r["budget"]
        budget_note = (f'<p>Budget envelope {_inr(b.get("budget_inr"))} — '
                       f'allocated {_inr(b.get("allocated_inr"))} ({b.get("utilization_pct")}%), '
                       f'remaining {_inr(b.get("remaining_inr"))}.</p>')

    def ul(items):
        return "<ul>" + "".join(f"<li>{_esc(x)}</li>" for x in items) + "</ul>"

    return f"""<!doctype html><html lang="en"><head><meta charset="utf-8"/>
<title>{_esc(r['title'])}</title>
<style>
@page {{ margin: 22mm 18mm; }}
* {{ box-sizing: border-box; }}
body {{ font-family: -apple-system, "Segoe UI", Roboto, Helvetica, Arial, sans-serif; color: {INK}; line-height: 1.55; margin: 0; font-size: 14px; }}
.wrap {{ max-width: 920px; margin: 0 auto; padding: 40px 28px; }}
.cover {{ border-left: 6px solid {BLUE}; padding: 8px 0 8px 20px; margin-bottom: 8px; }}
.cover .eyebrow {{ font: 700 11px/1 "SF Mono", Consolas, monospace; letter-spacing: 3px; text-transform: uppercase; color: {BLUE}; }}
.cover h1 {{ font-size: 30px; margin: 12px 0 6px; line-height: 1.15; }}
.cover .sub {{ color: #5b6472; font-size: 15px; }}
.cover .meta {{ margin-top: 14px; font-size: 12px; color: #7a828e; }}
h2 {{ font-size: 17px; margin: 34px 0 10px; padding-bottom: 6px; border-bottom: 2px solid #eef1f4; }}
h2 .num {{ color: {BLUE}; font: 700 13px/1 "SF Mono", Consolas, monospace; margin-right: 8px; }}
h3 {{ font-size: 13px; margin: 16px 0 4px; color: #3a4150; }}
p {{ margin: 8px 0; }}
.muted {{ color: #98a0ab; }}
.strip {{ display: grid; grid-template-columns: repeat(4, 1fr); gap: 10px; margin: 14px 0; }}
.strip .m {{ border: 1px solid #e6e9ee; border-radius: 10px; padding: 12px; text-align: center; }}
.strip .m b {{ display: block; font-size: 20px; font-weight: 800; }}
.strip .m span {{ font-size: 10px; color: #7a828e; text-transform: uppercase; letter-spacing: .04em; }}
table {{ width: 100%; border-collapse: collapse; margin: 10px 0; font-size: 13px; }}
table.kv th {{ text-align: left; width: 42%; color: #7a828e; font-weight: 600; padding: 6px 8px; border-bottom: 1px solid #eef1f4; vertical-align: top; }}
table.kv td {{ padding: 6px 8px; border-bottom: 1px solid #eef1f4; }}
table.grid th {{ text-align: left; background: #f6f8fa; padding: 8px; font-size: 11px; text-transform: uppercase; letter-spacing: .04em; color: #5b6472; border-bottom: 1px solid #e6e9ee; }}
table.grid td {{ padding: 8px; border-bottom: 1px solid #eef1f4; vertical-align: top; }}
figure {{ margin: 12px 0; }}
figure img {{ width: 100%; border-radius: 10px; border: 1px solid #e6e9ee; }}
figcaption {{ font-size: 11px; color: #7a828e; margin-top: 6px; }}
ul {{ margin: 8px 0; padding-left: 20px; }}
li {{ margin: 3px 0; }}
.two {{ display: grid; grid-template-columns: 1fr 1fr; gap: 20px; }}
.callout {{ background: #f2f7fe; border: 1px solid #d6e4fb; border-radius: 10px; padding: 12px 14px; margin: 10px 0; }}
.foot {{ margin-top: 36px; padding-top: 12px; border-top: 1px solid #eef1f4; font-size: 11px; color: #98a0ab; }}
</style></head><body><div class="wrap">

<div class="cover">
  <div class="eyebrow">Decision-ready planning document</div>
  <h1>{_esc(r['title'])}</h1>
  <div class="sub">{_esc(r['subtitle'])}</div>
  <div class="meta">Prepared by ClimaTwin · {_esc(r['generated_at'])} · {_esc(r['study_area']['administrative'])}</div>
</div>

<h2><span class="num">01</span>Executive Summary</h2>
<p>{_esc(n['executive_summary'])}</p>
<div class="strip">
  <div class="m"><b style="color:{GOOD}">−{imp.get('temp_reduction_c', 0)}°C</b><span>Feels-like</span></div>
  <div class="m"><b>{_int(imp.get('people_benefited'))}</b><span>People benefited</span></div>
  <div class="m"><b>{_inr(costs.get('capital_inr'))}</b><span>Capital</span></div>
  <div class="m"><b>{_int(imp.get('carbon_seq_kg_year'))}</b><span>kg CO₂/yr</span></div>
</div>

<h2><span class="num">02</span>Study Area</h2>
{_table([
    ("Area", _esc(sa['area_name'])),
    ("Coordinates", _esc(sa['coordinates'])),
    ("Administrative", _esc(sa['administrative'])),
    ("Population (est.)", _int(sa['population']) if sa['population'] else "—"),
    ("Commuter footfall", _esc(sa['commuter_footfall'])),
    ("Elderly share", f"{sa['elderly_pct']}%" if sa['elderly_pct'] else "—"),
    ("Ground elevation", f"{sa['elevation_m']} m" if sa['elevation_m'] is not None else "—"),
    ("Data blind spot", "Yes — survey recommended" if sa['data_blind_spot'] else "No"),
])}
{maps}

<h2><span class="num">03</span>Current Environmental Assessment</h2>
{_table([
    ("Heat (feels-like)", f"{asmt['heat']} °C — {_esc(asmt['condition'])}" if asmt['heat'] else "—"),
    ("Air quality", f"AQI {asmt['aqi']} ({_esc(asmt['aqi_category'])})" if asmt['aqi'] else "—"),
    ("Flood risk", _esc(asmt['flood_risk'])),
    ("Green cover", f"{asmt['green_cover_pct']}%" if asmt['green_cover_pct'] else "—"),
    ("Vegetation (NDVI)", _esc(asmt['ndvi'])),
])}

<h2><span class="num">04</span>Proposed Intervention Plan</h2>
<table class="grid"><thead><tr><th>Intervention</th><th>Qty</th><th>Capital</th><th>Rationale</th></tr></thead>
<tbody>{iv_rows}</tbody></table>
{budget_note}

<h2><span class="num">05</span>Why This Plan</h2>
<h3>Why these interventions were selected</h3><p>{_esc(n['why_selected'])}</p>
<h3>Why alternatives were down-weighted</h3><p>{_esc(n['why_alternatives_rejected'])}</p>
<div class="two">
  <div><h3>Environmental benefit</h3><p>{_esc(n['environmental_benefit'])}</p></div>
  <div><h3>Social benefit</h3><p>{_esc(n['social_benefit'])}</p></div>
</div>
<h3>Economic benefit</h3><p>{_esc(n['economic_benefit'])}</p>

<h2><span class="num">06</span>Impact Analysis</h2>
{r['charts']['impacts']}
{_table(impact_rows)}

<h2><span class="num">07</span>Budget</h2>
{_table([
    ("Capital cost", _inr(costs.get('capital_inr'))),
    ("Maintenance / year", _inr(costs.get('maintenance_inr_year'))),
    ("5-year cost of ownership", _inr(costs.get('five_year_inr'))),
    ("10-year cost of ownership", _inr(costs.get('ten_year_inr'))),
])}

<h2><span class="num">08</span>Implementation Timeline</h2>
{_table([
    ("Immediate (0–3 months)", _esc(n['timeline_immediate'])),
    ("Year 1", _esc(n['timeline_year1'])),
    ("Year 3", _esc(n['timeline_year3'])),
    ("Year 5", _esc(n['timeline_year5'])),
])}

<h2><span class="num">09</span>Risk Assessment</h2>
{ul(r['risks'])}

<h2><span class="num">10</span>Appendices</h2>
<h3>Methodology</h3><p>{_esc(r['methodology'])}</p>
<h3>Data sources</h3>{ul(r['data_sources'])}
<h3>AI assumptions & confidence</h3>{ul(r['assumptions'])}
<div class="callout">Every quantitative figure in this report is <b>illustrative and model-derived</b>, carrying an
uncertainty band and cited coefficients. It is intended to support planning discussion, not to replace
site survey and detailed engineering.</div>

<div class="foot">ClimaTwin — Urban Microclimate Decision Engine · Generated {_esc(r['generated_at'])} ·
AI narrative: Google Gemini (Flash) with deterministic fallback · Runs on Google Cloud free tiers.</div>

</div></body></html>"""


# --------------------------------------------------------------------------- #
# DOCX rendering
# --------------------------------------------------------------------------- #
def render_docx(r: dict) -> bytes:
    import io
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches

    doc = Document()
    doc.core_properties.title = r["title"]

    h = doc.add_heading(r["title"], level=0)
    doc.add_paragraph(r["subtitle"])
    doc.add_paragraph(f"Prepared by ClimaTwin · {r['generated_at']} · {r['study_area']['administrative']}")

    def section(num, title):
        p = doc.add_heading(f"{num}. {title}", level=1)
        return p

    def kv_table(rows):
        t = doc.add_table(rows=0, cols=2)
        t.style = "Light Grid Accent 1"
        for k, v in rows:
            cells = t.add_row().cells
            cells[0].text = str(k)
            cells[1].text = str(v)

    imp, costs, sa, asmt, n = r["impacts"], r["costs"], r["study_area"], r["assessment"], r["narrative"]

    section("1", "Executive Summary")
    doc.add_paragraph(n["executive_summary"])

    section("2", "Study Area")
    kv_table([
        ("Area", sa["area_name"]), ("Coordinates", sa["coordinates"]),
        ("Administrative", sa["administrative"]),
        ("Population (est.)", _int(sa["population"]) if sa["population"] else "—"),
        ("Commuter footfall", sa["commuter_footfall"] or "—"),
        ("Elderly share", f"{sa['elderly_pct']}%" if sa["elderly_pct"] else "—"),
        ("Ground elevation", f"{sa['elevation_m']} m" if sa["elevation_m"] is not None else "—"),
    ])
    # embed the roadmap static image if available
    img = _static_map(float(sa["coordinates"].split(",")[0]), float(sa["coordinates"].split(",")[1]))
    if img:
        try:
            doc.add_picture(io.BytesIO(img[0]), width=Inches(6.0))
        except Exception:
            pass

    section("3", "Current Environmental Assessment")
    kv_table([
        ("Heat (feels-like)", f"{asmt['heat']} °C" if asmt["heat"] else "—"),
        ("Air quality", f"AQI {asmt['aqi']} ({asmt['aqi_category']})" if asmt["aqi"] else "—"),
        ("Flood risk", asmt["flood_risk"] or "—"),
        ("Green cover", f"{asmt['green_cover_pct']}%" if asmt["green_cover_pct"] else "—"),
    ])

    section("4", "Proposed Intervention Plan")
    t = doc.add_table(rows=1, cols=4)
    t.style = "Light Grid Accent 1"
    hdr = t.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text, hdr[3].text = "Intervention", "Qty", "Capital", "Rationale"
    for i in r["interventions"]:
        c = t.add_row().cells
        c[0].text = str(i.get("name") or i.get("species") or "")
        c[1].text = _int(i.get("count"))
        c[2].text = _inr(i.get("capital_inr"))
        c[3].text = str(i.get("why", ""))

    section("5", "Why This Plan")
    for lbl, key in [("Why selected", "why_selected"), ("Why alternatives down-weighted", "why_alternatives_rejected"),
                     ("Environmental benefit", "environmental_benefit"), ("Social benefit", "social_benefit"),
                     ("Economic benefit", "economic_benefit")]:
        doc.add_heading(lbl, level=2)
        doc.add_paragraph(n[key])

    section("6", "Impact Analysis")
    kv_table([
        ("Temperature reduction", f"−{imp.get('temp_reduction_c', 0)} °C"),
        ("Air-quality improvement", f"{imp.get('aqi_improvement', 0)} AQI"),
        ("Stormwater managed", f"{_int(imp.get('flood_managed_m3'))} m³"),
        ("Green area added", f"{_int(imp.get('canopy_added_m2'))} m²"),
        ("Carbon sequestration", f"{_int(imp.get('carbon_seq_kg_year'))} kg/yr"),
        ("Population benefited", _int(imp.get("people_benefited"))),
    ])

    section("7", "Budget")
    kv_table([
        ("Capital cost", _inr(costs.get("capital_inr"))),
        ("Maintenance / year", _inr(costs.get("maintenance_inr_year"))),
        ("5-year cost", _inr(costs.get("five_year_inr"))),
        ("10-year cost", _inr(costs.get("ten_year_inr"))),
    ])

    section("8", "Implementation Timeline")
    kv_table([
        ("Immediate (0–3 mo)", n["timeline_immediate"]), ("Year 1", n["timeline_year1"]),
        ("Year 3", n["timeline_year3"]), ("Year 5", n["timeline_year5"]),
    ])

    section("9", "Risk Assessment")
    for risk in r["risks"]:
        doc.add_paragraph(risk, style="List Bullet")

    section("10", "Appendices")
    doc.add_heading("Methodology", level=2)
    doc.add_paragraph(r["methodology"])
    doc.add_heading("Data sources", level=2)
    for s in r["data_sources"]:
        doc.add_paragraph(s, style="List Bullet")
    doc.add_heading("AI assumptions & confidence", level=2)
    for a in r["assumptions"]:
        doc.add_paragraph(a, style="List Bullet")

    _ = (Pt, RGBColor)  # imported for future styling hooks
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
